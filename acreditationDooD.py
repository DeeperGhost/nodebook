import os
import datetime

import shutil
import qrcode

import PyPDF2

from shutil import ignore_patterns
# import qrcode.image.svg

from docx import Document
from docx.shared import Inches


def count(dir, counter=0):
    pout = open('out.txt', 'w')
    "returns number of files in dir and subdirs"
    for pack in os.walk(dir):
        for f in pack[2]:
            pout.write(str(dir) +str(pack[0])+ str(f) +'\n' )
            counter += 1
    pout.close()
    return dir + " : " + str(counter) + "files"
# вывод списка файлов в папке
def dirFolder():
    dir = "M:/"
    # dir = "d:/s_hare"
    report = "D:\\report/report.txt"
    # report = "M://"
    # p = os.listdir(dir)
    # for i in p:
    #     print(i)
    with open(report, "w", encoding="utf-8") as filewrite:
        for r, d, f in os.walk(dir):
            filewrite.write(r+'\n')
            for i in f:
                filewrite.write(i+'\n')
    filewrite.close()


# пробегает по директории и создает файл с рестром папок с пометкой пуста папка или нет
def compil_ot_file(dirName,fileOutName):
    dirName1 = 'O:\\БАЗА УП НА АККРЕДИТАЦИЮ/'
    # dirName = 'D:\\up base'

    listName = ['аОПОП','аРПУД','прил 0 опоп','прил 1 КУГ','прил 2 УП','прил 3 МК','прил 4 РПУДы','прил 5 ПП НИР',
                'прил 6 пр ГИА','прил 7 ССК','прил 8 ЭБС','прил 9 МТО','прил 10 РОП']
    numb = 0
    numberPapok = 0

    bufSTR ="1"
    bufSTR2 = '2'

    fileOut = dirName1 + '\Struct\\'+fileOutName
    print(fileOut)
    with open(fileOut, "w", encoding="utf-8") as filewrite:
        for r, d, f in os.walk(dirName):
            for i in listName:
                # if(i in str(r)):
                if(str(r).endswith(i)):
                    # отрезает корень каталога из начала строки , Добавляет в конец имя нужной папки через ;
                    # r1 = r.replace(';','')
                    b = r[dirName.__len__():str(r).rfind(str(i))]+';'+str(i)
                    # print(b)
                    bufSTR = b + ';' + ' 0' + '\n'

                    # bufSTR = str(r)+'/'+ '0' +'\n'
                    filewrite.write(bufSTR)
                    for file in f:
                        # print(str(r)+'/'+str(file) +'\n')
                        # Исключаем скрытый системный файл , вопрос почему переодически крашится от вывода на печать имя фйлов
                        if(str(file)!='Thumbs.db') and (str(file).endswith('.pdf')):
                            # r1 = r.replace(';','')
                            # отрезает корень каталога из начала строки , Добавляет в конец имя нужной папки через ;
                            b = r[dirName.__len__():str(r).rfind(str(i))]+';'+str(i)
                            bufSTR = b + ';' + ' 1' + '\n'

                            if(bufSTR != bufSTR2):
                                filewrite.write(bufSTR)
                                numberPapok += 1
                                bufSTR2 = bufSTR
            numb += 1
            print("пройдено = " + str(numb) + ' найдено папок = '+str(numberPapok))
    print('процент заполнения = ' + str(100*numberPapok/numb)[0:4] + '%')

    # запись логов обновления
    logupdate = open(dirName1 + '\Struct\\'+'logupdate.txt','a')
    run = datetime.datetime.today()
    logupdate.write(str(run.__format__('%d-%m-%Y %H:%M:%S'))+';'+str(numb)+';'+ str(numberPapok)+';'+str(100*numberPapok/numb)[0:4]+'\n')
    logupdate.close()



    fout = open(dirName1 + '\Struct\\' + 'src.txt', 'w')
    fout.write(datetime.datetime.today().strftime("%d/%m/%Y:%H-%M-%S"))
    fout.close()

# Создание QR кода
def qrcodeprint():
    # Create qr code instance
    qr = qrcode.QRCode(
        version = 1,
        error_correction = qrcode.constants.ERROR_CORRECT_H,
        box_size = 10,
        border = 4,
    )

    # The data that you want to store
    data = "https://www.dvfu.ru/about/rectorate/4915/the-department-of-organization-of-educational-activities/"
    probel = "\n"
    name = "Департамент организации образовательной деятельсти , - За подмену листа, убью сука!"
    # Add data
    qr.add_data(data)
    qr.add_data(probel)
    qr.add_data(name)
    qr.make(fit=True)

    # Create an image from the QR Code instance
    img = qr.make_image()

    # Save it somewhere, change the extension as needed:
    # img.save("image.png")
    # img.save("image.bmp")
    # img.save("image.jpeg")
    img.save("image.jpg")

#копирует директорию с pdf
def copy_to_pdf(directory,directoryToPdf,school):
    print(school)
    # рабочая схема
    shutil.copytree(directory+school,directoryToPdf+'/'+school,ignore=ignore_patterns('*.doc', '*.docx','*.txt','*.pli','*.xml','*.plx','*.plm','*.db',
                                                                    '*.xlsx','*.doc','*.xls','*.xlsm','*.rtf','*Арсеньев*',
                                                                    '*аспирантура*','*Большой Камень*','*Дальнегорск*',
                                                                    '*Находка*','*Уссурийск*','*ЦРПДО*',
                                                                    '*Электронные версии ОС ВО ДВФУ*',
                                                                    '*Электронные версии ФГОС ВО*',
                                                                    '*Электронные версии ФГОС ВО 3++*','*УВЦ*',
                                                                    '*российско-австралийская*','*российско-американская*',
                                                                    '*е выходя*','*2018 г.н. 08.04.01*','*2018 г.н. 15.04.01*',
                                                                    '*Б1.Б.1_Философские проблемы науки и техники-1-2.pdf*',
                                                                    '*Б1.Б.5_Общая теория динамических систем1-1-2*'))
    print("Закопировано")

#копирует директорию с pdf ОПОП
def copy_to_pdfOPOP(directory,directoryToPdf,school):
    print(school)
    # рабочая схема
    shutil.copytree(directory+school,directoryToPdf+'/'+school,ignore=ignore_patterns('*.doc', '*.docx','*.txt','*.pli','*.xml','*.plx','*.plm','*.db',
                                                                                      '*.xlsx','*.doc','*.xls','*.xlsm','*.rtf','*Арсеньев*',
                                                                                      '*аспирантура*','*Большой Камень*','*Дальнегорск*',
                                                                                      '*Находка*','*Уссурийск*','*ЦРПДО*',
                                                                                      '*Электронные версии ОС ВО ДВФУ*',
                                                                                      '*Электронные версии ФГОС ВО*',
                                                                                      '*Электронные версии ФГОС ВО 3++*','*УВЦ*',
                                                                                      '*российско-австралийская*','*российско-американская*',
                                                                                      '*е выходя*','*2018 г.н. 08.04.01*','*2018 г.н. 15.04.01*',
                                                                                      '*Б1.Б.1_Философские проблемы науки и техники-1-2.pdf*',
                                                                                      '*Б1.Б.5_Общая теория динамических систем1-1-2*',
                                                                                      '*аРПУД*','*прил 0*','*прил 1*','*прил 2*','*прил 3*','*прил 4*','*прил 5*',
                                                                                      '*прил 6*','*прил 7*','*прил 8*','*прил 9*','*прил 10*'))
    print("Закопировано")

#поиск в пдф
def findPDFTEST(pathtest):
    print('start \n')
    # file = 'pdftest.pdf'
    file = 'sample.pdf'
    pl = open(pathtest+'\\'+file, 'rb')
    plread = PyPDF2.PdfFileReader(pl)
    test = plread.getPage(0)
    print(test.extractText().encode('utf-8'))
    # for i in testencript:
    #     print(i)
    # # text20 = getpage20.extractText()
    # print(plread.numPages)
    # print(getpage20.encode('utf-8'))
    print('finish \n')





# выполняет задание по обновлению
def updateQuest():
    print("updateQuest")
    t= 'O:\БАЗА УП НА АККРЕДИТАЦИЮ\ШЭМ\Бакалавриат\\2014 г.н. ЗФО 38.03.01 Бухгалтерский учет, анализ и аудит\аОПОП'
    t1= '\ШЭМ\Бакалавриат\\2014 г.н. ЗФО 38.03.01 Бухгалтерский учет, анализ и аудит\аОПОП'
    dirToPDF = 'D:\\up base PDF\\UPDATE'
    dir1 = "O:\\БАЗА УП НА АККРЕДИТАЦИЮ"
    fileQuest = "ОБНОВЛЕНИЕ.txt"

    logUpdate = open('D:\\up base PDF\\logUpdate.txt', 'a')
    f = open(dir1+'\\'+fileQuest,'r')
    print(dir1+'\\'+fileQuest)
    for line in f:
        print(line)
    # shutil.copy(dir1+fileQuest,r'D:/1.txt')
        str1 = 'M'+line[1:]
        str1 = str1.replace('\n','')
        print(str1)
        str2 = str1.replace('M:\\БАЗА ОПОП на 2019-2020 уч.г',dirToPDF)

        run = datetime.datetime.today()
        logUpdate.write(str(run.__format__('%d-%m-%Y %H:%M:%S'))+' ' + str2+'\n')

        shutil.copytree(str1,str2,ignore=ignore_patterns('*.doc', '*.docx','*.txt','*.pli','*.xml','*.plx','*.plm','*.db',
                                                                                          '*.xlsx','*.doc','*.xls','*.xlsm','*.rtf','*Арсеньев*',
                                                                                          '*аспирантура*','*Большой Камень*','*Дальнегорск*',
                                                                                          '*Находка*','*Уссурийск*','*ЦРПДО*',
                                                                                          '*Электронные версии ОС ВО ДВФУ*',
                                                                                          '*Электронные версии ФГОС ВО*',
                                                                                          '*Электронные версии ФГОС ВО 3++*','*УВЦ*',
                                                                                          '*российско-австралийская*','*российско-американская*',
                                                                                          '*е выходя*'))


    logUpdate.close()

def copySHEN(directory,directoryToPdf,school):
    listName = ['аОПОП','аРПУД','прил 0 опоп','прил 1 КУГ','прил 2 УП','прил 3 МК','прил 4 РПУДы','прил 5 ПП НИР',
                'прил 6 пр ГИА','прил 7 ССК','прил 8 ЭБС','прил 9 МТО','прил 10 РОП']

    shutil.copytree(directory+school,directoryToPdf+'/'+school,ignore=ignore_patterns('*аОПОП*','*аРПУД*','*прил 0 опоп*','*прил 3 МК*','*прил 4 РПУДы*','*прил 5 ПП НИР*',
                                                                                      '*прил 6 пр ГИА*','*прил 7 ССК*','*прил 8 ЭБС*','*прил 9 МТО*','*прил 10 РОП*'))


def copyDrozdova(directory,directoryToPdf,school):
    listName = ['аОПОП','аРПУД','прил 0 опоп','прил 1 КУГ','прил 2 УП','прил 3 МК','прил 4 РПУДы','прил 5 ПП НИР',
                'прил 6 пр ГИА','прил 7 ССК','прил 8 ЭБС','прил 9 МТО','прил 10 РОП']

    shutil.copytree(directory,directoryToPdf+'/'+school,ignore=ignore_patterns('*прил 1 КУГ*','*прил 2 УП*'))

#     Гененрирует ШАблон DOCX
def TemplateDOCX():
    print("go gog")
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('first item in ordered list', style='List Number')

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = ((3, '101', 'Spam'),(7, '422', 'Eggs'),(4, '631', 'Spam, spam, eggs, and spam'))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


# main функция
def accreditationCopy():
    dirName = 'M:\\БАЗА ОПОП на 2019-2020 уч.г/'
    # dirName = 'D:\\up base'
    # dirToPDF = 'D:\\up base PDF'
    dirToPDF = 'W:\\Осеев Е.Е/up base PDF'

    # fileOutAll = 'file.txt'
    # fileOutPDF = 'filePDF.txt'

    SCHOOL = ['ВИ-ШРМИ','ИШ','ШБМ','ШЕН','ШИГН','ШЦЭ','ШЭМ','ЮШ','филиал в г. Арсеньеве',
              'филиал в г. Большой Камень','филиал в г. Находке','филиал в г. Уссурийске (ШП)']


    # копирует пдфки
    for i in SCHOOL:
        copy_to_pdf(dirName,dirToPDF,i)

def accreditation():
    # dirName = 'O:\\БАЗА УП НА АККРЕДИТАЦИЮ/'
    # dirName = 'M:\\БАЗА ОПОП на 2019-2020 уч.г/' #19 год набора
    dirName = 'M:\\БАЗА ОПОП 2020 г.н/'
    # dirName = 'https://bb.dvfu.ru/bbcswebdav/orgs/FUDOOD/19'
    # dirName = 'D:\\up base'
    dirToPDF = 'D:\\up base PDF'
    dirToPDF_OPOP = 'D:\\up base PDF/opop'
    # dirToPDF = 'W:\\Осеев Е.Е/up base PDF'
    SHENDOOD = 'D:\\SHENDOOD'
    SHENDROZDOVA = 'W:\ШЕН_Дроздова'
    fileOutAll = 'file.txt'
    fileOutPDF = 'filePDF.txt'

    SCHOOL = ['ВИ-ШРМИ','ИШ','ШБМ','ШЕН','ШИГН','ШЦЭ','ШЭМ','ЮШ','филиал в г. Арсеньеве',
              'филиал в г. Большой Камень','филиал в г. Находке','филиал в г. Уссурийске (ШП)']

    # Создать общий выходной файл
    compil_ot_file(dirName,fileOutAll)
    # TemplateDOCX()
    # Создать выходной файл для PDF директории
    # compil_ot_file(dirToPDF, fileOutPDF)

    # qrcodeprint()

    # копирует пдфки
    # copyUGPI()
    # for i in SCHOOL:
    #     copy_to_pdf(dirName,dirToPDF,i)

    # Копирует ОПЫ ОПЫ
    # for i in SCHOOL:
    #     copy_to_pdfOPOP(dirName,dirToPDF_OPOP,i)
    # findPDFTEST(pathtest=dirToPDF)
    #


    # copy_to_pdf(dirName,dirToPDF,'ШИГН') #ШИГН1
    # copy_to_pdf(dirName,dirToPDF,'ЮШ') #ЮШ
    # copy_to_pdf(dirName,dirToPDF,'ШЕН')
    # copy_to_pdf(dirName,dirToPDF,'ИШ')
    # copy_to_pdf(dirName,dirToPDF,'ВИ-ШРМИ')

    # copy_to_pdf(dirName,dirToPDF,'ШЭМ')
    # copy_to_pdf(dirName,dirToPDF,'ШБМ')
    # copy_to_pdf(dirName,dirToPDF,'ШЦЭ')

    # обновление по запросу
    # updateQuest()

    # Копирует УПЫ И КУГИ с ДООД
    # copySHEN(dirName,SHENDOOD,'ШЕН')

    # Копирует из дроздова не УПЫ и КУГИ
    # copyDrozdova(SHENDROZDOVA,SHENDOOD,'ШЕНДРОЗДОВА')


